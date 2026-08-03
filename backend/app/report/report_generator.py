"""
Generador de Informe Pericial
================================

Convierte un FormalCaseReport (ver pipeline/formal_pipeline.py) en un
documento con dos secciones, tal como se define en el diseño de
MODEXRE:

  1. Resumen ejecutivo: lenguaje no técnico, apto para un juez o
     responsable no especializado en IA.
  2. Anexo técnico: metodología, resultados detallados por evento,
     explicación SHAP, y la cadena de custodia íntegra para su
     verificación independiente.

Este módulo produce Markdown (fuente única y legible), que se
convierte a .docx solo en el momento de generar el entregable final
para un caso concreto, no en cada ejecución de tests.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

from app.pipeline.formal_pipeline import FormalCaseReport


ATTACK_LABEL_MAP = {0: "Tráfico normal", 1: "Actividad potencialmente maliciosa"}


def _format_finding_summary(report: FormalCaseReport) -> tuple[int, int]:
    n_malicious = sum(1 for f in report.findings if f.prediction["label"] == 1)
    n_normal = len(report.findings) - n_malicious
    return n_malicious, n_normal


def generate_executive_summary(report: FormalCaseReport) -> str:
    n_malicious, n_normal = _format_finding_summary(report)
    total = len(report.findings)

    verification = report.custody_verification or {}
    integrity_line = (
        "La cadena de custodia del caso ha sido verificada matemáticamente "
        "y no presenta indicios de alteración."
        if verification.get("valid")
        else "**ADVERTENCIA: la verificación de la cadena de custodia ha detectado "
        "posibles inconsistencias. Ver Anexo Técnico, apartado de Cadena de Custodia.**"
    )

    lines = [
        f"# Informe Pericial — Caso {report.case_id}",
        "",
        f"*Generado el {datetime.now(timezone.utc).isoformat()}*",
        "",
        "## Resumen Ejecutivo",
        "",
        f"Se ha analizado el tráfico de red correspondiente a la fuente "
        f"`{report.source_name}` (fichero: `{report.file_path}`), aplicando un "
        f"sistema automatizado de detección de intrusiones basado en inteligencia "
        f"artificial (XGBoost), cuyas decisiones han sido interpretadas mediante "
        f"técnicas de IA explicable (SHAP).",
        "",
        f"Del total de {total} eventos analizados con evidencia suficiente para "
        f"su evaluación, **{n_malicious}** han sido clasificados como actividad "
        f"potencialmente maliciosa y **{n_normal}** como tráfico normal.",
        "",
        integrity_line,
        "",
        "Los detalles técnicos completos, incluyendo la justificación de cada "
        "clasificación y el registro íntegro de la cadena de custodia, se "
        "encuentran en el Anexo Técnico de este documento.",
    ]
    return "\n".join(lines)


def generate_technical_annex(report: FormalCaseReport) -> str:
    lines = ["", "## Anexo Técnico", "", "### Metodología", ""]
    lines += [
        "- **Fuente de datos:** " + report.source_name,
        "- **Normalización:** OCSF (Open Cybersecurity Schema Framework), "
        "clase Detection Finding (class_uid=2004).",
        "- **Modelo:** XGBoost, entrenado y congelado en modo Laboratorio; "
        "usado exclusivamente en modo inferencia en este caso.",
        "- **Explicabilidad:** SHAP (TreeExplainer), explicación local por evento.",
        "",
        "### Resultados por evento",
        "",
    ]

    for i, finding in enumerate(report.findings):
        pred = finding.prediction
        label_text = pred.get("attack_cat", ATTACK_LABEL_MAP.get(pred["label"], "Desconocido"))
        finding_info = finding.ocsf_event.get("finding_info", {})
        src = finding.ocsf_event.get("src_endpoint", {})
        dst = finding.ocsf_event.get("dst_endpoint", {})

        lines.append(f"#### Evento #{i + 1}")
        lines.append("")
        lines.append(f"- **Clasificación del modelo:** {label_text} "
                      f"(probabilidad: {pred['probability']:.3f}, modelo: {pred['model_version']})")
        if "class_probabilities" in pred:
            lines.append("- **Distribución de probabilidad por categoría:**")
            for cat, p in sorted(pred["class_probabilities"].items(), key=lambda kv: -kv[1]):
                lines.append(f"  - {cat}: {p:.3f}")
        if finding_info.get("title"):
            lines.append(f"- **Alerta original del IDS:** {finding_info['title']} "
                          f"(categoría: {', '.join(finding_info.get('types', [])) or 'n/d'})")
        lines.append(f"- **Origen → Destino:** {src.get('ip', 'n/d')}:{src.get('port', 'n/d')} "
                      f"→ {dst.get('ip', 'n/d')}:{dst.get('port', 'n/d')}")
        lines.append("- **Variables más influyentes en la decisión (SHAP):**")
        for feat in finding.explanation["top_features"]:
            lines.append(f"  - `{feat['feature']}`: {feat['shap_value']:+.4f}")
        lines.append("")

    lines.append("### Cadena de Custodia")
    lines.append("")
    verification = report.custody_verification or {}
    lines.append(f"- **Total de eslabones registrados:** {verification.get('total_records', 'n/d')}")
    lines.append(f"- **Íntegra:** {'Sí' if verification.get('valid') else 'NO — ver incidencias'}")
    if verification.get("issues"):
        lines.append("- **Incidencias detectadas:**")
        for issue in verification["issues"]:
            lines.append(f"  - {issue}")

    return "\n".join(lines)


def _finding_category(finding) -> str:
    """Nombre de categoría de un hallazgo, sea FindingResult (objeto)
    o el dict equivalente ya serializado (report_to_json)."""
    pred = finding.prediction if hasattr(finding, "prediction") else finding["prediction"]
    return pred.get("attack_cat") or ATTACK_LABEL_MAP.get(pred["label"], "Desconocido")


def group_findings_by_category(report: FormalCaseReport) -> dict[str, list]:
    """Agrupa los hallazgos de un caso por categoría de ataque
    detectada, preservando el orden de primera aparición de cada
    categoría en la evidencia (no alfabético, para que el informe
    refleje la secuencia real de los hechos)."""
    grouped: dict[str, list] = {}
    for finding in report.findings:
        cat = _finding_category(finding)
        grouped.setdefault(cat, []).append(finding)
    return grouped


def build_category_summary(report: FormalCaseReport) -> list[dict[str, Any]]:
    """Resumen por categoría, pensado para mostrarse en una interfaz
    ANTES de generar el informe final, de forma que el analista pueda
    decidir qué categorías incluir: nº de eventos, IPs de origen y
    destino implicadas, probabilidad media asignada por el modelo."""
    grouped = group_findings_by_category(report)
    summary = []
    for cat, findings in grouped.items():
        src_ips = sorted({f.ocsf_event.get("src_endpoint", {}).get("ip")
                           for f in findings if f.ocsf_event.get("src_endpoint", {}).get("ip")})
        dst_ips = sorted({f.ocsf_event.get("dst_endpoint", {}).get("ip")
                           for f in findings if f.ocsf_event.get("dst_endpoint", {}).get("ip")})
        probs = [f.prediction["probability"] for f in findings]
        summary.append({
            "categoria": cat,
            "n_eventos": len(findings),
            "src_ips": src_ips,
            "dst_ips": dst_ips,
            "probabilidad_media": sum(probs) / len(probs) if probs else 0.0,
        })
    return summary


def aggregate_shap_by_category(findings: list, top_n: int = 5) -> list[dict[str, Any]]:
    """Agrega las contribuciones SHAP de un grupo de hallazgos (todos
    de la misma categoría, típicamente) en un único ranking — la suma
    de |valor SHAP| de cada variable a través de todos los eventos.

    Es la misma lógica que ya usaba generate_report_docx_judicial
    internamente para su bloque 'Análisis Explicable (XAI) del
    Incidente'; se extrae aquí como función reutilizable para que la
    interfaz pueda mostrar el mismo resumen antes de generar el
    informe, sin tener que generar el .docx solo para verlo.
    """
    feature_totals: dict[str, float] = {}
    for finding in findings:
        explanation = finding.explanation if hasattr(finding, "explanation") else finding["explanation"]
        for feat in explanation["top_features"]:
            feature_totals[feat["feature"]] = feature_totals.get(feat["feature"], 0.0) + abs(feat["shap_value"])
    ranked = sorted(feature_totals.items(), key=lambda kv: -kv[1])[:top_n]
    return [{"feature": name, "contribucion_agregada": val} for name, val in ranked]


def filter_report_by_categories(report: FormalCaseReport, categories: list[str]) -> FormalCaseReport:
    """Devuelve una COPIA del informe con solo los hallazgos cuya
    categoría esté en `categories`. No modifica el informe original
    (así se puede volver a filtrar de otra forma sin rehacer todo el
    pipeline pericial ni tocar la cadena de custodia ya registrada).

    `categories` vacío o None equivale a "todas" (sin filtrar).
    """
    if not categories:
        selected_findings = list(report.findings)
    else:
        wanted = set(categories)
        selected_findings = [f for f in report.findings if _finding_category(f) in wanted]

    return FormalCaseReport(
        case_id=report.case_id,
        source_name=report.source_name,
        file_path=report.file_path,
        findings=selected_findings,
        custody_verification=report.custody_verification,
    )


def generate_report_markdown(report: FormalCaseReport) -> str:
    return generate_executive_summary(report) + "\n" + generate_technical_annex(report)


def report_to_json(report: FormalCaseReport, chain_records: list[Any]) -> dict:
    """Serializa un FormalCaseReport (+ los registros de su cadena de
    custodia) a un dict JSON-friendly, formato de entrada del
    generador de .docx (scripts/generate_report_docx.js).

    `chain_records` se recibe explícitamente (en vez de que este
    módulo consulte CustodyChain directamente) para no acoplar la
    capa de reporting a la capa de custodia más de lo necesario: quien
    llama ya tiene ambos objetos tras ejecutar FormalCaseRunner.
    """
    return {
        "case_id": report.case_id,
        "source_name": report.source_name,
        "file_path": report.file_path,
        "custody_verification": report.custody_verification,
        "findings": [
            {
                "prediction": f.prediction,
                "explanation": f.explanation,
                "finding_info": f.ocsf_event.get("finding_info", {}),
                "src_endpoint": f.ocsf_event.get("src_endpoint", {}),
                "dst_endpoint": f.ocsf_event.get("dst_endpoint", {}),
            }
            for f in report.findings
        ],
        "custody_chain": [
            {
                "index": r.index, "timestamp": r.timestamp, "operation": r.operation,
                "component": r.component, "record_hash": r.record_hash, "notes": r.notes,
            }
            for r in chain_records
        ],
    }


def generate_report_docx_via_node(report: FormalCaseReport, chain_records: list[Any], output_path: str) -> str:
    """[ALTERNATIVA, no usada por la app] Genera el .docx invocando el
    script Node (docx-js). Requiere Node.js + la librería npm 'docx'
    en el entorno — dependencia que NO se quiere imponer a quien
    ejecute MODEXRE en su máquina (ver generate_report_docx, que usa
    python-docx y no tiene esta dependencia). Se conserva por si en
    algún momento conviene regenerar el informe desde este mismo
    entorno de desarrollo con más control de estilo via docx-js.
    """
    import json
    import subprocess
    import tempfile

    data = report_to_json(report, chain_records)
    script_path = Path(__file__).parent / "scripts" / "generate_report_docx.js"

    with tempfile.NamedTemporaryFile("w", suffix=".json", delete=False, encoding="utf-8") as tmp:
        json.dump(data, tmp, ensure_ascii=False)
        tmp_json_path = tmp.name

    try:
        subprocess.run(
            ["node", str(script_path), tmp_json_path, output_path],
            check=True, capture_output=True, text=True,
        )
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Fallo al generar el .docx: {e.stderr}") from e
    finally:
        Path(tmp_json_path).unlink(missing_ok=True)

    return output_path


def _add_page_number_field(paragraph) -> None:
    """Inserta un campo de número de página dinámico (Word lo
    resuelve al abrir/imprimir), usando la API de bajo nivel de
    python-docx porque no hay wrapper de alto nivel para campos."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _add_toc_field(doc) -> None:
    """Inserta un índice dinámico (campo TOC). Word lo muestra vacío
    ('Haga clic para actualizar el índice') hasta que el usuario
    actualiza el campo (clic derecho > Actualizar campo, o F9) — es
    el mismo comportamiento que tiene cualquier tabla de contenido
    automática de Word, no un límite de este generador."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Haga clic derecho aquí y seleccione \"Actualizar campo\" para generar el índice."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def _kv_paragraph(doc, label: str, value: str):
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(str(value))
    return p


def _bullet_list(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _build_category_bar_chart(category_summary: list[dict[str, Any]]) -> Optional[bytes]:
    """Genera un gráfico de barras (nº de eventos por categoría) como
    PNG en memoria, para insertarlo en el informe técnico. Devuelve
    None si matplotlib no está disponible o no hay datos — el informe
    debe poder generarse igualmente sin la gráfica en ese caso."""
    if not category_summary:
        return None
    try:
        import io as _io
        import matplotlib
        matplotlib.use("Agg")  # backend sin interfaz gráfica, necesario en servidor
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    categorias = [row["categoria"] for row in category_summary]
    valores = [row["n_eventos"] for row in category_summary]
    colores = ["#2e7d32" if cat == "Normal" else "#c62828" for cat in categorias]

    fig, ax = plt.subplots(figsize=(8, 4))
    ax.bar(categorias, valores, color=colores)
    ax.set_ylabel("Número de eventos")
    ax.set_title("Distribución de eventos por categoría clasificada")
    for i, v in enumerate(valores):
        ax.text(i, v, str(v), ha="center", va="bottom")
    fig.tight_layout()

    buffer = _io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150)
    plt.close(fig)
    return buffer.getvalue()


def _apply_document_styles(doc, heading_font: str = "Arial") -> None:
    """Aplica fuente/tamaño/color a los estilos de encabezado del
    documento, calcados de los informes de referencia aportados:
    encabezados en `heading_font` a 20/18/15pt (Heading 1/2/3), color
    azul #2F5496 (el azul estándar de Word usado en ambos informes de
    referencia), cuerpo de texto en Calibri 11pt.
    """
    from docx.shared import Pt, RGBColor

    HEADING_BLUE = RGBColor(0x2F, 0x54, 0x96)

    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
    except KeyError:
        pass

    heading_sizes = {"Heading 1": 20, "Heading 2": 18, "Heading 3": 15, "Heading 4": 12}
    for style_name, size_pt in heading_sizes.items():
        try:
            style = doc.styles[style_name]
            style.font.name = heading_font
            style.font.size = Pt(size_pt)
            style.font.color.rgb = HEADING_BLUE
            style.font.bold = True
        except KeyError:
            continue


def generate_report_docx(
    report: FormalCaseReport,
    chain_records: list[Any],
    output_path: str | None = None,
    *,
    cliente: str = "No especificado",
    perito: str = "Sistema MODEXRE (análisis automatizado, supervisado por el operador del sistema)",
    referencia_caso: str | None = None,
) -> bytes:
    """Genera el informe pericial en .docx con python-docx (sin
    dependencias externas más allá de lo que ya instala
    requirements.txt).

    Estructura alineada con un informe pericial informático estándar
    (portada, declaración de cumplimiento, índice, resumen ejecutivo,
    descripción del caso, recolección de datos, metodología, análisis
    de la evidencia, resultados, conclusiones, recomendaciones,
    anexos), rellenada con datos REALES del caso allí donde MODEXRE
    los tiene (cadena de custodia, metodología del pipeline, hallazgos
    con SHAP) — solo `cliente` y `perito` son datos que el operador
    del sistema debe aportar, porque no son deducibles del análisis.

    Devuelve los bytes del documento. Si se proporciona `output_path`,
    además los escribe a disco en esa ruta.
    """
    import io
    from datetime import datetime, timezone
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION

    data = report_to_json(report, chain_records)
    verification = data["custody_verification"] or {}
    findings = data["findings"]
    n_malicious = sum(1 for f in findings if f["prediction"]["label"] != 0
                        and f["prediction"].get("attack_cat") != "Normal")
    n_normal = len(findings) - n_malicious
    fecha_hoy = datetime.now(timezone.utc).strftime("%d/%m/%Y")
    referencia_caso = referencia_caso or data["case_id"]

    # Categorías presentes en los hallazgos, para "Resultados"/"Conclusiones"
    categorias_presentes: dict[str, int] = {}
    for f in findings:
        cat = f["prediction"].get("attack_cat") or ATTACK_LABEL_MAP.get(f["prediction"]["label"], "Desconocido")
        categorias_presentes[cat] = categorias_presentes.get(cat, 0) + 1

    doc = Document()
    _apply_document_styles(doc, heading_font="Arial")
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        # Numeración de página en el pie
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number_field(footer_p)

    # ============================== PORTADA ==============================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run("\n\n\nInforme Pericial Informático").bold = True
    doc.paragraphs[-1].runs[0].font.size = Pt(28)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.add_run(f"Análisis de tráfico de red — Caso {referencia_caso}").font.size = Pt(16)

    doc.add_paragraph()
    for label, value in [
        ("Cliente / destinatario", cliente),
        ("Perito / responsable del análisis", perito),
        ("Fecha del informe", fecha_hoy),
        ("Sistema utilizado", "MODEXRE — Sistema de detección forense de intrusiones con IA explicable"),
        ("Fuente de la evidencia", data["source_name"]),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    doc.add_page_break()

    # ==================== DECLARACIÓN DE CUMPLIMIENTO ====================
    doc.add_heading("Declaración de Cumplimiento", level=1)
    doc.add_paragraph(
        f"El presente análisis se ha realizado mediante el sistema MODEXRE, "
        f"aplicando un pipeline automatizado y trazable de principio a fin: "
        f"ingesta de evidencia, normalización al estándar OCSF (Open "
        f"Cybersecurity Schema Framework), extracción de características, "
        f"clasificación mediante un modelo de aprendizaje automático (XGBoost) "
        f"previamente entrenado y certificado, y explicación de cada decisión "
        f"mediante técnicas de IA explicable (SHAP)."
    )
    doc.add_paragraph(
        "Cada paso del proceso ha quedado registrado en una cadena de custodia "
        "digital (hash-chain), que permite verificar matemáticamente que la "
        "evidencia y los resultados intermedios no han sido alterados entre "
        "el momento de la ingesta y la emisión de este informe. El detalle "
        "completo de dicha cadena se incluye en el Anexo de este documento."
    )
    doc.add_paragraph(
        "El diseño metodológico de este análisis se ha inspirado en los "
        "principios de la norma ISO/IEC 27037:2012 (identificación, "
        "recolección, adquisición y preservación de evidencia digital), "
        "ISO/IEC 27042:2015 (análisis e interpretación de evidencia digital) "
        "y la guía NIST SP 800-86 (Guide to Integrating Forensic Techniques "
        "into Incident Response)."
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "El modelo de clasificación empleado nunca se reentrena en modo de "
        "análisis pericial: se congela tras su validación en un entorno de "
        "laboratorio separado, y su integridad (hash del fichero de modelo) "
        "se verifica en el momento de cargarlo para este caso."
    )

    doc.add_page_break()

    # ============================== ÍNDICE ==============================
    doc.add_heading("Índice", level=1)
    _add_toc_field(doc)
    doc.add_page_break()

    # ========================= RESUMEN EJECUTIVO =========================
    doc.add_heading("Resumen Ejecutivo", level=1)
    doc.add_paragraph(
        f"Se ha analizado el tráfico de red correspondiente a la fuente "
        f"\"{data['source_name']}\" (fichero: \"{data['file_path']}\"), aplicando "
        f"un sistema automatizado de detección de intrusiones basado en "
        f"inteligencia artificial (XGBoost), cuyas decisiones han sido "
        f"interpretadas mediante técnicas de IA explicable (SHAP)."
    )
    p = doc.add_paragraph()
    p.add_run(f"Del total de {len(findings)} eventos analizados, ")
    p.add_run(str(n_malicious)).bold = True
    p.add_run(" se han clasificado como actividad potencialmente maliciosa y ")
    p.add_run(str(n_normal)).bold = True
    p.add_run(" como tráfico normal.")

    if categorias_presentes:
        p = doc.add_paragraph()
        p.add_run("Categorías detectadas: ").bold = True
        p.add_run(", ".join(f"{cat} ({n})" for cat, n in sorted(categorias_presentes.items(), key=lambda kv: -kv[1])))

    integrity_p = doc.add_paragraph()
    if verification.get("valid"):
        integrity_p.add_run(
            "La cadena de custodia del caso ha sido verificada matemáticamente "
            "y no presenta indicios de alteración."
        )
    else:
        integrity_p.add_run(
            "ADVERTENCIA: la verificación de la cadena de custodia ha detectado "
            "posibles inconsistencias. Ver apartado de Cadena de Custodia."
        ).bold = True

    doc.add_page_break()

    # ========================= DESCRIPCIÓN DEL CASO =========================
    doc.add_heading("1. Descripción del Caso", level=1)
    doc.add_paragraph(
        f"El presente informe se refiere al caso {referencia_caso}, que "
        f"consiste en el análisis de evidencia de red procedente de la fuente "
        f"\"{data['source_name']}\", con el objetivo de determinar la presencia "
        f"de actividad de reconocimiento de red (reconnaissance) o escaneo de "
        f"puertos (port scanning) mediante un sistema de clasificación basado "
        f"en inteligencia artificial."
    )

    doc.add_heading("Antecedentes", level=2)
    doc.add_paragraph(
        f"La evidencia analizada procede de {data['source_name']}, ingerida "
        f"en el sistema el {fecha_hoy}. El fichero de origen ({data['file_path']}) "
        f"ha sido tratado como evidencia digital desde el momento de su "
        f"ingesta, generando el primer eslabón de la cadena de custodia con su "
        f"hash criptográfico (SHA-256)."
    )

    doc.add_heading("Objetivos del Informe", level=2)
    _bullet_list(doc, [
        "Determinar si la evidencia analizada contiene indicios de actividad "
        "de reconocimiento de red o escaneo de puertos.",
        "Clasificar cada evento de la evidencia mediante un modelo de IA "
        "entrenado y certificado previamente, con explicación local (SHAP) "
        "de cada decisión.",
        "Preservar la integridad y trazabilidad de la evidencia durante todo "
        "el proceso, mediante una cadena de custodia digital verificable.",
        "Presentar los resultados en un formato accesible tanto para un "
        "lector no técnico (Resumen Ejecutivo) como para un perito o "
        "auditor técnico (Anexo Técnico).",
    ])

    doc.add_heading("Alcance del Informe", level=2)
    _bullet_list(doc, [
        f"Fuente de evidencia analizada: {data['source_name']}.",
        f"Modelo de clasificación utilizado: {findings[0]['prediction']['model_version'] if findings else 'n/d'}.",
        "El análisis se limita a los eventos contenidos en el fichero de "
        "evidencia proporcionado; no incluye la investigación de sistemas o "
        "dispositivos no representados en dicha evidencia.",
        "El sistema clasifica y explica; no ejecuta ninguna acción de "
        "respuesta (bloqueo, aislamiento, etc.) sobre la infraestructura "
        "analizada.",
    ])

    doc.add_heading("Consideraciones Legales", level=2)
    doc.add_paragraph(
        "Este análisis se ha realizado sobre evidencia ya en posesión del "
        "solicitante del informe. La cadena de custodia digital implementada "
        "documenta cada transformación aplicada a la evidencia (ingesta, "
        "normalización, agregación, extracción de características, "
        "clasificación), permitiendo su verificación independiente. Se "
        "recomienda que un perito legal evalúe la admisibilidad de este "
        "informe conforme a la jurisdicción aplicable."
    )

    doc.add_page_break()

    # ========================= RECOLECCIÓN DE DATOS =========================
    doc.add_heading("2. Recolección de Datos", level=1)

    doc.add_heading("Tipos de Datos", level=2)
    doc.add_paragraph(
        f"La evidencia corresponde a eventos de tipo \"{data['source_name']}\", "
        f"normalizados al estándar OCSF (clase Detection Finding, "
        f"category_uid=2, class_uid=2004) tras su ingesta."
    )

    doc.add_heading("Métodos de Recolección", level=2)
    doc.add_paragraph(
        f"La evidencia se ha ingerido mediante el conector correspondiente a "
        f"\"{data['source_name']}\" del sistema MODEXRE, que calcula el hash "
        f"SHA-256 del fichero completo en el momento de la ingesta, antes de "
        f"cualquier transformación."
    )

    doc.add_heading("Cadena de Custodia", level=2)
    _kv_paragraph(doc, "Total de eslabones registrados", verification.get("total_records", "n/d"))
    _kv_paragraph(doc, "Integridad verificada", "Sí" if verification.get("valid") else "NO — ver incidencias")
    doc.add_paragraph(
        "Cada eslabón registra: la operación realizada, el componente de "
        "software responsable (con versión), el hash de entrada y de salida "
        "de esa operación, y el hash del eslabón anterior, formando una "
        "cadena encadenada verificable. El detalle completo se incluye en el "
        "Anexo I de este documento."
    )

    doc.add_heading("Descripción de la Evidencia", level=2)
    _kv_paragraph(doc, "Fichero de origen", data["file_path"])
    _kv_paragraph(doc, "Fuente", data["source_name"])
    _kv_paragraph(doc, "Fecha de ingesta", fecha_hoy)

    doc.add_heading("Registro de Procedimientos", level=2)
    doc.add_paragraph(
        "El siguiente listado resume, en orden cronológico, cada operación "
        "aplicada a la evidencia (detalle completo en el Anexo I):"
    )
    for rec in data["custody_chain"]:
        doc.add_paragraph(f"{rec['operation']} — {rec['notes'] or 'sin notas adicionales'}", style="List Number")

    doc.add_page_break()

    # ============================= METODOLOGÍA =============================
    doc.add_heading("3. Metodología", level=1)

    doc.add_heading("Descripción de los Métodos Utilizados", level=2)
    doc.add_paragraph(
        "El análisis sigue un pipeline de cinco etapas, cada una registrada "
        "en la cadena de custodia: (1) ingesta de la evidencia bruta; (2) "
        "normalización al estándar OCSF; (3) agregación causal por origen y "
        "ventana temporal, para detectar patrones de comportamiento que un "
        "único evento aislado no revela (p. ej., cuántos puertos distintos ha "
        "tocado un mismo origen); (4) extracción de un vector de "
        "características numéricas por evento; y (5) clasificación mediante "
        "un modelo XGBoost, con explicación local de cada decisión mediante "
        "SHAP."
    )

    doc.add_heading("Herramientas Forenses Utilizadas", level=2)
    _bullet_list(doc, [
        "MODEXRE (sistema propio) — orquestación del pipeline y cadena de custodia.",
        "OCSF (Open Cybersecurity Schema Framework) — normalización de eventos.",
        "XGBoost — modelo de clasificación.",
        "SHAP (TreeExplainer) — explicabilidad local de cada decisión del modelo.",
    ])

    doc.add_heading("Procesos de Análisis", level=2)
    doc.add_paragraph(
        f"El modelo utilizado en este caso "
        f"({findings[0]['prediction']['model_version'] if findings else 'n/d'}) "
        f"fue entrenado y validado en modo Laboratorio, con datos separados de "
        f"los analizados en este informe, y congelado (verificación de "
        f"integridad por hash) antes de su uso aquí. El modelo NO se ha "
        f"reentrenado ni ajustado con la evidencia de este caso."
    )

    doc.add_heading("Validación de Resultados", level=2)
    doc.add_paragraph(
        "Cada clasificación se acompaña de su probabilidad y, cuando el "
        "modelo es multi-clase, de la distribución de probabilidad completa "
        "sobre todas las categorías conocidas por el modelo — no solo el "
        "veredicto más probable. Además, cada decisión incluye las variables "
        "que más influyeron en ella (SHAP), permitiendo auditar el "
        "razonamiento del modelo caso por caso."
    )

    doc.add_heading("Consideraciones Éticas y Legales", level=2)
    doc.add_paragraph(
        "El sistema no ejecuta ninguna acción automática sobre la "
        "infraestructura analizada (bloqueo, aislamiento, notificación). "
        "Su función se limita a clasificar y explicar; cualquier acción "
        "derivada de este informe es responsabilidad del destinatario."
    )

    doc.add_page_break()

    # ======================= ANÁLISIS DE LA EVIDENCIA =======================
    doc.add_heading("4. Análisis de la Evidencia", level=1)

    category_summary = build_category_summary(report)
    grouped_findings = group_findings_by_category(report)

    doc.add_paragraph(
        f"Se han analizado {len(findings)} eventos, agrupados en "
        f"{len(category_summary)} categoría(s) según la clasificación del "
        f"modelo. Se presenta a continuación un resumen analítico por "
        f"categoría — no un desglose evento a evento, dado el volumen "
        f"habitual de eventos en un caso real — seguido de una "
        f"representación gráfica de la distribución."
    )

    doc.add_heading("Tabla de Resultados Agrupados", level=2)
    results_table = doc.add_table(rows=1, cols=5)
    results_table.style = "Light Grid Accent 1"
    hdr = results_table.rows[0].cells
    for idx, title in enumerate(["Categoría", "Nº eventos", "Prob. media", "IPs origen", "Variable SHAP principal"]):
        hdr[idx].text = title
    for row_data in category_summary:
        cat = row_data["categoria"]
        top_shap = aggregate_shap_by_category(grouped_findings[cat], top_n=1)
        top_feat_text = top_shap[0]["feature"] if top_shap else "n/d"

        row = results_table.add_row().cells
        row[0].text = cat
        row[1].text = str(row_data["n_eventos"])
        row[2].text = f"{row_data['probabilidad_media']:.3f}"
        row[3].text = ", ".join(row_data["src_ips"][:5]) + ("..." if len(row_data["src_ips"]) > 5 else "") or "n/d"
        row[4].text = top_feat_text
    doc.add_paragraph()

    doc.add_heading("Distribución de Eventos por Categoría", level=2)
    chart_bytes = _build_category_bar_chart(category_summary)
    if chart_bytes:
        doc.add_picture(io.BytesIO(chart_bytes), width=Cm(15))
    doc.add_paragraph()

    doc.add_heading("Variables Más Influyentes por Categoría (SHAP agregado)", level=2)
    doc.add_paragraph(
        "Para cada categoría, se muestran las variables cuya contribución "
        "SHAP acumulada (valor absoluto, sumado sobre todos los eventos de "
        "esa categoría) ha sido mayor — es decir, las que más han pesado en "
        "el conjunto de decisiones del modelo para ese tipo de evento."
    )
    for row_data in category_summary:
        cat = row_data["categoria"]
        doc.add_heading(cat, level=3)
        shap_table = doc.add_table(rows=1, cols=2)
        shap_table.style = "Light Grid Accent 1"
        hdr = shap_table.rows[0].cells
        hdr[0].text, hdr[1].text = "Variable", "Contribución agregada (|SHAP|)"
        for feat in aggregate_shap_by_category(grouped_findings[cat], top_n=5):
            row = shap_table.add_row().cells
            row[0].text = feat["feature"]
            row[1].text = f"{feat['contribucion_agregada']:.3f}"
        doc.add_paragraph()


    doc.add_page_break()

    # ============================== RESULTADOS ==============================
    doc.add_heading("5. Resultados", level=1)

    doc.add_heading("Organización de los Resultados", level=2)
    doc.add_paragraph(
        f"De los {len(findings)} eventos analizados, la distribución por "
        f"categoría clasificada por el modelo es la siguiente:"
    )
    result_table = doc.add_table(rows=1, cols=2)
    result_table.style = "Light Grid Accent 1"
    hdr = result_table.rows[0].cells
    hdr[0].text, hdr[1].text = "Categoría", "Nº de eventos"
    for cat, n in sorted(categorias_presentes.items(), key=lambda kv: -kv[1]):
        row = result_table.add_row().cells
        row[0].text = cat
        row[1].text = str(n)

    doc.add_heading("Interpretación de los Resultados", level=2)
    if n_malicious > 0:
        doc.add_paragraph(
            f"La presencia de {n_malicious} evento(s) clasificado(s) como "
            f"actividad potencialmente maliciosa, junto con las variables "
            f"identificadas por SHAP en cada caso (particularmente las "
            f"relativas a la diversidad de puertos/hosts destino tocados por "
            f"un mismo origen, cuando están disponibles), es consistente con "
            f"un patrón de reconocimiento de red o escaneo de puertos."
        )
    else:
        doc.add_paragraph(
            "No se han identificado eventos clasificados como actividad "
            "maliciosa en la evidencia analizada."
        )

    doc.add_heading("Limitaciones", level=2)
    doc.add_paragraph(
        "Los resultados dependen de la cobertura de características que la "
        "fuente de evidencia puede proporcionar. Fuentes con menor riqueza de "
        "datos de flujo (p. ej., alertas de IDS que no incluyen todas las "
        "estadísticas que sí ofrece un dataset de investigación como "
        "CICFlowMeter) ofrecen una clasificación con menor cantidad de señal "
        "disponible para el modelo, aunque el sistema está diseñado para "
        "tolerar esta situación sin fallar."
    )

    doc.add_page_break()

    # ============================== CONCLUSIONES ==============================
    doc.add_heading("6. Conclusiones", level=1)
    doc.add_paragraph(
        f"Sobre la base del análisis realizado, de los {len(findings)} eventos "
        f"examinados, {n_malicious} presentan indicios de actividad "
        f"potencialmente maliciosa según el modelo de clasificación empleado, "
        f"con la justificación local (SHAP) documentada para cada uno en el "
        f"apartado de Análisis de la Evidencia."
    )
    doc.add_paragraph(
        "La cadena de custodia digital de este caso ha sido verificada "
        "matemáticamente" + (" y no presenta indicios de alteración."
        if verification.get("valid") else ", presentando incidencias que se detallan en el Anexo I.")
    )
    doc.add_paragraph(
        "Los resultados presentados se basan exclusivamente en el análisis "
        "automatizado de la evidencia proporcionada, interpretado mediante "
        "técnicas de IA explicable, y no constituyen por sí solos una "
        "determinación legal de responsabilidad."
    )

    doc.add_page_break()

    # ============================= RECOMENDACIONES =============================
    doc.add_heading("7. Recomendaciones", level=1)
    if n_malicious > 0:
        origenes = sorted({
            f["src_endpoint"].get("ip") for f in findings
            if f["prediction"]["label"] != 0 and f["src_endpoint"].get("ip")
        })
        _bullet_list(doc, [
            f"Revisar y, si procede, restringir el acceso desde el/los origen(es) "
            f"identificado(s) como maliciosos: {', '.join(origenes) if origenes else 'ver detalle por evento'}.",
            "Correlacionar estos hallazgos con los registros del propio "
            "IDS/IPS/firewall de origen para una verificación cruzada adicional.",
            "Si se confirma la actividad de reconocimiento, evaluar el "
            "endurecimiento de la exposición de los puertos/servicios tocados "
            "por el origen identificado.",
            "Conservar este informe y el fichero de evidencia original junto "
            "con la base de datos de la cadena de custodia, para su eventual "
            "verificación independiente.",
        ])
    else:
        doc.add_paragraph(
            "No se han identificado acciones correctivas específicas a partir "
            "de este análisis. Se recomienda mantener la monitorización "
            "continua de la fuente analizada."
        )

    doc.add_page_break()

    # ================================ ANEXOS ================================
    doc.add_heading("Anexo I — Cadena de Custodia Completa", level=1)
    chain_table = doc.add_table(rows=1, cols=4)
    chain_table.style = "Light Grid Accent 1"
    hdr = chain_table.rows[0].cells
    for idx, title in enumerate(["#", "Operación", "Componente", "Hash (16 primeros caracteres)"]):
        hdr[idx].text = title
    for rec in data["custody_chain"]:
        row = chain_table.add_row().cells
        row[0].text = str(rec["index"])
        row[1].text = rec["operation"]
        row[2].text = rec["component"]
        row[3].text = rec["record_hash"][:16] + "..."

    if verification.get("issues"):
        doc.add_heading("Incidencias Detectadas en la Verificación", level=2)
        _bullet_list(doc, verification["issues"])

    buffer = io.BytesIO()
    doc.save(buffer)
    doc_bytes = buffer.getvalue()

    if output_path:
        Path(output_path).write_bytes(doc_bytes)

    return doc_bytes


# =========================================================================
# INFORME PERICIAL JUDICIAL (formato de presentación ante un Juzgado)
# =========================================================================
#
# A diferencia de generate_report_docx (informe técnico completo, pensado
# para un auditor o perito contrario que necesita revisar cada detalle),
# este generador produce el documento en el formato que realmente se
# presenta ante una autoridad judicial: cabecera de juzgado con folios,
# base legal, un bloque "CASO FORENSE N.º X" por cada tipología de ataque
# detectada en la evidencia (no solo un listado plano de eventos),
# conclusión pericial global, admisibilidad, recomendaciones, y las
# declaraciones finales de objetividad/independencia con firma del perito.
#
# Estructura y registro basados en el ejemplo real de informe pericial
# aportado para el proyecto (Diligencias Previas, Juzgado de Instrucción),
# adaptados para generarse dinámicamente a partir de los datos reales de
# cualquier FormalCaseReport de MODEXRE.

CATEGORY_TECHNICAL_DESCRIPTIONS: dict[str, dict[str, str]] = {
    "PortScan": {
        "nombre_caso": "ESCANEO DE PUERTOS (PORT SCAN)",
        "caracteristicas": (
            "Desde el punto de vista técnico, los ataques de Port Scan se "
            "caracterizan por la realización de múltiples intentos de conexión "
            "hacia distintos puertos de destino en un intervalo de tiempo "
            "reducido. Esta actividad suele presentar patrones claramente "
            "diferenciables del tráfico legítimo: elevado número de conexiones "
            "a distintos puertos desde una misma dirección origen, ausencia de "
            "sesiones prolongadas o intercambio significativo de datos, alta "
            "rotación de puertos en intervalos temporales muy cortos, y tasas "
            "de fallo de conexión superiores a las esperadas en escenarios "
            "normales. Estos elementos apuntan a un comportamiento "
            "automatizado orientado a la enumeración de servicios, más que al "
            "uso funcional de los mismos."
        ),
        "evidencia_generica": (
            "La evidencia digital recopilada para este caso incluye múltiples "
            "flujos de red con características homogéneas, que, analizados de "
            "forma agregada, permiten reconstruir el comportamiento típico de "
            "un ataque de Port Scan: distribuciones de puertos de destino "
            "altamente dispersas, paquetes de tamaño reducido y escaso "
            "intercambio bidireccional, y secuencias temporales regulares que "
            "indican el uso de herramientas automatizadas. A diferencia de los "
            "ataques DoS o de Fuerza Bruta, el impacto inmediato de un Port "
            "Scan sobre la disponibilidad del sistema puede ser limitado; no "
            "obstante, su detección resulta crítica, ya que suele preceder a "
            "ataques más graves y constituye una fase clara de preparación."
        ),
    },
    "DDoS": {
        "nombre_caso": "DENEGACIÓN DE SERVICIO DISTRIBUIDA (DDoS)",
        "caracteristicas": (
            "Los ataques de Denegación de Servicio Distribuida (DDoS) se "
            "caracterizan por un volumen de tráfico anormalmente elevado "
            "dirigido hacia un mismo destino, con el objetivo de agotar sus "
            "recursos de red o de cómputo. Los indicadores técnicos observados "
            "incluyen: tasas de paquetes y bytes muy superiores a la línea "
            "base del tráfico legítimo, duración de flujo reducida y "
            "repetitiva, y patrones de origen consistentes con tráfico "
            "generado por herramientas automatizadas o botnets."
        ),
        "evidencia_generica": (
            "La evidencia digital muestra un volumen de conexiones y de "
            "tráfico incompatible con un uso legítimo del servicio, "
            "concentrado en ventanas temporales cortas y con un patrón "
            "repetitivo característico de tráfico generado por herramientas "
            "automatizadas."
        ),
    },
    "DoS": {
        "nombre_caso": "DENEGACIÓN DE SERVICIO (DoS)",
        "caracteristicas": (
            "Los ataques de Denegación de Servicio (DoS) se caracterizan por "
            "un volumen o frecuencia de peticiones dirigidas a un servicio "
            "concreto con la finalidad de agotar su capacidad de respuesta. "
            "Se observan indicadores como duración de flujo anómala, un "
            "número de paquetes por conexión muy superior al habitual, y una "
            "concentración de la actividad en un intervalo temporal reducido."
        ),
        "evidencia_generica": (
            "Los registros de tráfico examinados muestran un patrón de "
            "peticiones incompatible con el uso funcional normal del "
            "servicio, con una intensidad y regularidad que apuntan a un "
            "origen automatizado."
        ),
    },
    "BruteForce": {
        "nombre_caso": "FUERZA BRUTA SOBRE SERVICIOS DE AUTENTICACIÓN",
        "caracteristicas": (
            "Los ataques de fuerza bruta se caracterizan por múltiples "
            "intentos sucesivos de autenticación contra un mismo servicio, "
            "en intervalos de tiempo muy cortos. Los indicadores técnicos "
            "observados incluyen: elevada frecuencia de intentos de conexión, "
            "duración de sesión característicamente corta por intento, y "
            "concentración de la actividad sobre un puerto o servicio "
            "concreto de autenticación."
        ),
        "evidencia_generica": (
            "La evidencia digital muestra un patrón de intentos de conexión "
            "repetidos y sistemáticos, incompatible con el comportamiento "
            "habitual de un usuario legítimo, y compatible con el uso de "
            "herramientas automatizadas de prueba de credenciales."
        ),
    },
    "Reconnaissance": {
        "nombre_caso": "RECONOCIMIENTO DE RED (RECONNAISSANCE)",
        "caracteristicas": (
            "Las actividades de reconocimiento se caracterizan por la "
            "exploración sistemática de la infraestructura de red objetivo, "
            "con el fin de identificar sistemas activos, servicios expuestos "
            "y posibles vulnerabilidades, sin explotar activamente ninguna de "
            "ellas en esta fase."
        ),
        "evidencia_generica": (
            "La evidencia digital es compatible con una fase de exploración "
            "previa a un ataque, mostrando un patrón de acceso sistemático a "
            "distintos recursos de la red objetivo."
        ),
    },
}

_DEFAULT_CATEGORY_DESCRIPTION = {
    "nombre_caso": "ACTIVIDAD ANÓMALA DE RED",
    "caracteristicas": (
        "El tráfico examinado presenta características estadísticas "
        "significativamente distintas del patrón de tráfico legítimo "
        "observado como referencia, según el modelo de clasificación "
        "empleado."
    ),
    "evidencia_generica": (
        "La evidencia digital recopilada muestra un conjunto de flujos de "
        "red cuyas características agregadas son coherentes entre sí y "
        "distintas del comportamiento normal de la red analizada."
    ),
}


def _judicial_footer_text(ref: str, folio: int, total_folios: int, caso_titulo: str | None = None) -> list[str]:
    lines = [f"Ref.: {ref} — Folio {folio} de {total_folios}", "JUZGADO DE INSTRUCCIÓN"]
    if caso_titulo:
        lines.append(caso_titulo)
    return lines


def generate_report_docx_judicial(
    report: FormalCaseReport,
    chain_records: list[Any],
    output_path: str | None = None,
    *,
    juzgado: str = "JUZGADO DE INSTRUCCIÓN N.º [A COMPLETAR]",
    procedimiento: str = "DILIGENCIAS PREVIAS N.º [A COMPLETAR] / [AÑO]",
    ref_expediente: str | None = None,
    perito_nombre: str = "[NOMBRE DEL PERITO]",
    perito_titulacion: str = "Analista forense en ciberseguridad — Especialista en Inteligencia Artificial",
    perito_colegiado: str = "[N.º de colegiado, si aplica]",
    cliente: str = "No especificado",
) -> bytes:
    """Genera el informe pericial en el formato de presentación judicial:
    cabecera de juzgado con folios, base legal, un bloque "CASO FORENSE"
    por cada tipología de ataque detectada en la evidencia, conclusión
    pericial global, admisibilidad, recomendaciones, y declaraciones
    finales de objetividad/independencia con firma.

    A diferencia de generate_report_docx (informe técnico completo con
    el detalle de cada evento individual), este agrupa los hallazgos POR
    TIPOLOGÍA DE ATAQUE detectada, en línea con cómo se presenta la
    prueba pericial ante una autoridad judicial: no interesa el detalle
    de cada paquete, sino la valoración pericial de cada patrón de
    ataque identificado.

    Los campos `juzgado`, `procedimiento`, `ref_expediente`, `perito_*`
    son datos que debe aportar el operador del sistema para un caso
    real — MODEXRE no puede deducirlos de la evidencia analizada.
    """
    import io
    from datetime import datetime, timezone
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    data = report_to_json(report, chain_records)
    verification = data["custody_verification"] or {}
    findings = data["findings"]
    fecha_hoy = datetime.now(timezone.utc).strftime("%d/%m/%Y")
    ref_expediente = ref_expediente or f"MODEXRE/{data['case_id']}"

    # Agrupar hallazgos por categoría detectada (excluyendo Normal), en
    # el orden en que aparecen por primera vez en la evidencia.
    categorias_orden: list[str] = []
    hallazgos_por_categoria: dict[str, list[dict]] = {}
    for f in findings:
        cat = f["prediction"].get("attack_cat") or ATTACK_LABEL_MAP.get(f["prediction"]["label"], "Desconocido")
        if cat == "Normal":
            continue
        if cat not in hallazgos_por_categoria:
            hallazgos_por_categoria[cat] = []
            categorias_orden.append(cat)
        hallazgos_por_categoria[cat].append(f)

    doc = Document()
    _apply_document_styles(doc, heading_font="Calibri")
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    def add_footer(caso_titulo: str | None = None):
        """Pie de página con referencia de expediente + número de
        folio dinámico (campo PAGE de Word)."""
        footer_p = doc.sections[-1].footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.add_run(f"Ref.: {ref_expediente} — Folio ").bold = True
        _add_page_number_field(footer_p)
        footer_p.add_run(" de ")
        # NUMPAGES: total de páginas del documento (campo dinámico)
        run = footer_p.add_run()
        fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "NUMPAGES"
        fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_sep); run._r.append(fld_end)
        if caso_titulo:
            doc.add_paragraph()  # separación visual mínima ya la da el propio pie

    add_footer()

    # ==================== IDENTIFICACIÓN DEL EXPEDIENTE ====================
    # Bloque compacto (sin página de portada separada): conserva los
    # datos identificativos del caso al principio del documento.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(juzgado).bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("INFORME PERICIAL FORENSE DIGITAL SOBRE INCIDENTE DE RED "
                "DETECTADO MEDIANTE INTELIGENCIA ARTIFICIAL EXPLICABLE (XAI)").bold = True

    doc.add_paragraph()
    for label, value in [
        ("Ref.", ref_expediente),
        ("Procedimiento", procedimiento),
        ("Emitido por", f"{perito_nombre} — {perito_titulacion}"),
        ("N.º de colegiado", perito_colegiado),
        ("A requerimiento de", cliente),
        ("Fecha de emisión", fecha_hoy),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("*** DOCUMENTO CONFIDENCIAL *** — ")
    run.bold = True
    p.add_run(
        "Su difusión queda limitada conforme a la normativa procesal "
        "aplicable en materia de protección de datos y confidencialidad "
        "de las actuaciones judiciales."
    ).italic = True

    doc.add_paragraph()

    # ============================ 1. OBJETO ============================
    doc.add_heading("1. Objeto del Informe", level=1)
    doc.add_paragraph(
        "El presente informe pericial tiene por objeto el análisis técnico "
        "y forense de una serie de eventos de tráfico de red que presentan "
        "características compatibles con distintos tipos de ataques "
        "informáticos, detectados mediante el uso combinado de un modelo "
        "de clasificación de aprendizaje automático (XGBoost) y técnicas "
        "de inteligencia artificial explicable (SHAP)."
    )
    doc.add_paragraph(
        "El dictamen se emite con la finalidad de determinar la existencia "
        "de indicios técnicos razonables de actividad maliciosa, así como "
        "de justificar de forma comprensible, reproducible y trazable las "
        "conclusiones alcanzadas, de modo que puedan ser evaluadas por "
        "terceros, incluidos órganos judiciales, auditores técnicos o "
        "peritos independientes."
    )
    doc.add_paragraph(
        "El análisis realizado no se limita a la obtención de una "
        "clasificación automática (tráfico legítimo frente a tráfico "
        "malicioso), sino que incorpora una interpretación de los factores "
        "técnicos que han conducido a dicha clasificación. Las "
        "conclusiones se expresan en términos de probabilidad técnica y "
        "coherencia con patrones conocidos de ataque, evitando "
        "afirmaciones categóricas que excedan el alcance del análisis "
        "realizado."
    )
    p = doc.add_paragraph()
    p.add_run("El informe NO incluye: ").bold = True
    p.add_run(
        "identificación de autoría humana del ataque; análisis del "
        "impacto económico; investigación sobre infraestructura física "
        "del sistema origen. El alcance queda limitado estrictamente al "
        "análisis de red y explicación técnica de las predicciones "
        "generadas por el modelo de IA empleado."
    )

    # ================ 2. IDENTIFICACIÓN DEL PERITO ================
    doc.add_heading("2. Identificación del Perito y Alcance de su Intervención", level=1)
    doc.add_paragraph(
        f"El presente informe ha sido elaborado por {perito_nombre}, en su "
        f"condición de {perito_titulacion}, actuando con plena autonomía "
        f"técnica y sin relación de dependencia con las partes implicadas "
        f"en el procedimiento."
    )
    doc.add_paragraph(
        "La intervención del perito se ha limitado al análisis técnico de "
        "la evidencia digital proporcionada, mediante el sistema MODEXRE, "
        "sin acceso ni intervención sobre los sistemas de producción "
        "origen de dicha evidencia."
    )

    # ============ 3. BASE LEGAL Y NORMATIVA APLICABLE ============
    doc.add_heading("3. Base Legal y Normativa Aplicable", level=1)
    doc.add_paragraph(
        "El presente informe se elabora conforme a los principios de "
        "objetividad, imparcialidad y neutralidad técnica exigidos a la "
        "prueba pericial. El diseño metodológico del análisis se ha "
        "inspirado en los principios de las normas ISO/IEC 27037:2012 "
        "(identificación, recolección, adquisición y preservación de "
        "evidencia digital), ISO/IEC 27042:2015 (análisis e "
        "interpretación de evidencia digital), y la guía NIST SP 800-86 "
        "(Guide to Integrating Forensic Techniques into Incident "
        "Response)."
    )

    # ============ 4. DOCUMENTACIÓN RECIBIDA ============
    doc.add_heading("4. Documentación y Evidencia Recibida", level=1)
    _kv_paragraph(doc, "Fuente de la evidencia", data["source_name"])
    _kv_paragraph(doc, "Fichero analizado", data["file_path"])
    _kv_paragraph(doc, "Fecha de recepción/ingesta", fecha_hoy)
    _kv_paragraph(doc, "Hash de integridad de la cadena de custodia",
                   "Verificado" if verification.get("valid") else "CON INCIDENCIAS — ver Anexo")

    doc.add_page_break()

    # ============ 5. METODOLOGÍA GENERAL ============
    doc.add_heading("5. Metodología General de Análisis", level=1)
    doc.add_paragraph(
        "El análisis se ha desarrollado mediante un proceso automatizado "
        "y trazable de cinco etapas, cada una registrada de forma "
        "verificable en una cadena de custodia digital: (1) ingesta de la "
        "evidencia bruta, con cálculo de su huella criptográfica (SHA-256); "
        "(2) normalización de los eventos al estándar internacional OCSF "
        "(Open Cybersecurity Schema Framework); (3) agregación causal por "
        "origen y ventana temporal, para identificar patrones de "
        "comportamiento no visibles en un único evento aislado; (4) "
        "extracción de un vector de características numéricas por evento; "
        "y (5) clasificación mediante un modelo de aprendizaje automático "
        "(XGBoost), con explicación local de cada decisión mediante la "
        "técnica SHAP (SHapley Additive exPlanations)."
    )

    doc.add_heading("Uso de Inteligencia Artificial Explicable (XAI)", level=2)
    doc.add_paragraph(
        "El uso de técnicas de IA explicable (XAI) resulta central en la "
        "metodología aplicada: no basta con obtener una clasificación "
        "automática, sino que resulta necesario justificar de forma "
        "comprensible y reproducible qué variables técnicas concretas han "
        "motivado dicha clasificación. Para ello se emplea SHAP, técnica "
        "que asigna a cada variable del modelo una contribución "
        "cuantificada a la decisión final, permitiendo reconstruir el "
        "razonamiento del modelo caso por caso."
    )

    doc.add_heading("Limitaciones Generales del Análisis", level=2)
    doc.add_paragraph(
        "El presente análisis está sujeto a las siguientes limitaciones, "
        "que deben tenerse en cuenta en su valoración: el modelo emplea un "
        "subconjunto de características derivables de la fuente de "
        "evidencia analizada, no la totalidad de las variables que podría "
        "ofrecer un dataset de investigación con instrumentación completa; "
        "las conclusiones se expresan en términos de probabilidad técnica, "
        "no de certeza absoluta; y el sistema no determina intencionalidad "
        "ni autoría humana, limitándose al análisis objetivo del tráfico "
        "de red."
    )

    doc.add_heading("Modelo de Detección Empleado", level=2)
    modelo_version = findings[0]["prediction"]["model_version"] if findings else "n/d"
    _kv_paragraph(doc, "Modelo", f"XGBoost — versión certificada {modelo_version}")
    doc.add_paragraph(
        "El modelo fue entrenado y validado en un entorno de laboratorio "
        "independiente, con datos distintos de la evidencia analizada en "
        "este informe, y posteriormente congelado (verificación de "
        "integridad mediante hash criptográfico del propio fichero de "
        "modelo) para su uso exclusivo en modo de inferencia. El modelo NO "
        "se ha reentrenado ni ajustado con la evidencia de este caso, "
        "garantizando que el sistema aplicado es exactamente el mismo que "
        "fue validado con anterioridad."
    )

    doc.add_page_break()

    # ============ BLOQUES POR CASO FORENSE (por categoría detectada) ============
    if not categorias_orden:
        doc.add_heading("6. Análisis de la Evidencia", level=1)
        doc.add_paragraph(
            "No se han identificado patrones de tráfico compatibles con "
            "ninguna tipología de ataque conocida en la evidencia analizada."
        )
    else:
        for case_num, categoria in enumerate(categorias_orden, start=1):
            hallazgos_cat = hallazgos_por_categoria[categoria]
            desc = CATEGORY_TECHNICAL_DESCRIPTIONS.get(categoria, _DEFAULT_CATEGORY_DESCRIPTION)
            caso_titulo = f"CASO FORENSE N.º {case_num} — ATAQUE DE {desc['nombre_caso']}"

            doc.add_heading(caso_titulo, level=1)

            doc.add_heading("Identificación del Incidente Analizado", level=2)
            origenes = sorted({f["src_endpoint"].get("ip") for f in hallazgos_cat if f["src_endpoint"].get("ip")})
            doc.add_paragraph(
                f"Se han identificado {len(hallazgos_cat)} evento(s) de la evidencia "
                f"analizada clasificados por el modelo como compatibles con un patrón "
                f"de {desc['nombre_caso'].lower()}"
                + (f", con origen en: {', '.join(origenes)}." if origenes else ".")
            )

            doc.add_heading("Características Técnicas del Incidente", level=2)
            doc.add_paragraph(desc["caracteristicas"])

            doc.add_heading("Evidencia Digital Asociada al Incidente", level=2)
            doc.add_paragraph(desc["evidencia_generica"])
            probs = [f["prediction"]["probability"] for f in hallazgos_cat]
            prob_media = sum(probs) / len(probs) if probs else 0.0
            doc.add_paragraph(
                f"La probabilidad media asignada por el modelo a los eventos de "
                f"este caso es de {prob_media:.3f} sobre 1.0."
            )

            doc.add_heading("Análisis Explicable (XAI) del Incidente", level=2)
            # Agregación simple de las variables más citadas en SHAP entre
            # los eventos de esta categoría, como aproximación al "análisis
            # global" del ejemplo de referencia.
            top_global = [
                (item["feature"], item["contribucion_agregada"])
                for item in aggregate_shap_by_category(hallazgos_cat, top_n=5)
            ]
            doc.add_paragraph(
                "Con el fin de interpretar las decisiones del modelo, se ha "
                "aplicado la técnica SHAP a cada evento individual de este caso. "
                "Las variables con mayor influencia agregada en la "
                "clasificación de estos eventos han sido:"
            )
            _bullet_list(doc, [f"{name} (contribución agregada: {val:.3f})" for name, val in top_global])
            doc.add_paragraph(
                "Esta explicación permite vincular directamente las métricas "
                "técnicas del tráfico analizado con el comportamiento esperado "
                "de esta tipología de ataque, proporcionando una justificación "
                "clara y reproducible de la clasificación."
            )

            doc.add_heading("Valoración Pericial del Caso", level=2)
            doc.add_paragraph(
                f"A la vista de los elementos analizados, se concluye que los "
                f"registros examinados presentan una coherencia técnica elevada "
                f"con un ataque de {desc['nombre_caso'].lower()}, en consonancia "
                f"con los patrones descritos en la literatura de seguridad "
                f"informática y análisis forense de red."
            )
            doc.add_paragraph(
                "Esta conclusión se formula en términos de probabilidad técnica "
                "elevada, sin atribuir intencionalidad subjetiva ni "
                "responsabilidad directa, limitándose estrictamente al análisis "
                "objetivo del tráfico de red."
            )
            doc.add_paragraph(f"Con ello se da por concluido el análisis correspondiente al CASO {case_num}.")
            doc.add_page_break()

    # ============ CONCLUSIÓN PERICIAL GLOBAL ============
    doc.add_heading("Conclusión Pericial Global del Informe", level=1)
    if categorias_orden:
        doc.add_paragraph(
            f"El análisis técnico y forense realizado ha permitido identificar "
            f"y evaluar, con criterios objetivos y reproducibles, "
            f"{len(categorias_orden)} tipología(s) de ataque diferenciada(s) en "
            f"la evidencia analizada: {', '.join(CATEGORY_TECHNICAL_DESCRIPTIONS.get(c, _DEFAULT_CATEGORY_DESCRIPTION)['nombre_caso'] for c in categorias_orden)}."
        )
    doc.add_paragraph(
        "Las conclusiones alcanzadas no se basan exclusivamente en una "
        "clasificación automática, sino en un análisis combinado que "
        "integra: los resultados de un modelo de detección entrenado y "
        "validado; la evidencia digital observable en los flujos de red "
        "analizados; y su interpretación mediante técnicas explicables "
        "(SHAP). Este enfoque permite afirmar que las inferencias "
        "realizadas se apoyan en criterios técnicos verificables, evitando "
        "decisiones opacas o arbitrarias."
    )
    doc.add_paragraph(
        "Desde el punto de vista probatorio, el sistema aplicado cumple "
        "con los principios fundamentales exigidos en el ámbito forense: "
        "reproducibilidad, trazabilidad, integridad y explicabilidad. "
        "Puede afirmarse que el sistema descrito constituye una "
        "herramienta válida de apoyo al análisis forense digital, siempre "
        "que sea utilizada bajo supervisión humana y con pleno "
        "conocimiento de sus limitaciones. Su aplicación no sustituye al "
        "criterio del perito, pero sí refuerza su capacidad analítica y la "
        "solidez técnica de las conclusiones alcanzadas."
    )

    doc.add_heading("Resultado Pericial", level=2)
    doc.add_paragraph("El perito concluye que:")
    _bullet_list(doc, [
        ("En el/los caso(s) analizado(s) concurren indicios técnicos sólidos, "
         "repetitivos y coherentes con ataques cibernéticos reales."
         if categorias_orden else
         "No se han identificado indicios técnicos de actividad maliciosa en "
         "la evidencia analizada."),
        "Las explicaciones generadas mediante IA explicable (XAI) son "
        "reproducibles, auditables y metodológicamente válidas como apoyo a "
        "la prueba pericial.",
        f"La evidencia digital se ha mantenido íntegra bajo cadena de "
        f"custodia controlada"
        + (" y verificada." if verification.get("valid") else ", si bien se han detectado incidencias — ver Anexo."),
    ])

    doc.add_heading("Admisibilidad y Aptitud para Juicio", level=2)
    doc.add_paragraph(
        "El informe cumple con los requisitos de objetividad, claridad y "
        "precisión exigibles a la prueba pericial, así como con los "
        "estándares internacionales de investigación digital referidos en "
        "el apartado de Base Legal. Por ello, este informe puede ser "
        "presentado ante la autoridad judicial competente como prueba "
        "pericial digital."
    )

    doc.add_page_break()

    # ============ RECOMENDACIONES ============
    doc.add_heading("Recomendaciones a la Autoridad Judicial", level=1)
    _bullet_list(doc, [
        "Admitir el presente informe como prueba pericial digital complementaria.",
        "Solicitar, si se considera oportuno, contrapericia independiente "
        "para validar la reproducibilidad del análisis.",
        "Relevar al perito para ampliación oral durante vista o juicio, si "
        "resultara necesario.",
    ])

    doc.add_heading("Recomendaciones Técnicas (para la entidad afectada)", level=1)
    recomendaciones_tecnicas = [
        "Implementar un IDS/IPS con capacidad de detección basada en firmas "
        "y en comportamiento.",
        "Registrar de forma sistemática las métricas de tráfico relevantes "
        "empleadas en este análisis (duración, paquetes y bytes por "
        "dirección, puertos/hosts distintos contactados por origen).",
        "Incorporar herramientas de IA explicable en los procesos internos "
        "de auditoría y en el SOC.",
        "Mantener actualización continua de los sistemas críticos "
        "expuestos.",
    ]
    if "PortScan" in categorias_orden or "Reconnaissance" in categorias_orden:
        recomendaciones_tecnicas.append(
            "Revisar y, si procede, restringir el acceso desde los orígenes "
            "identificados con actividad de escaneo en el presente informe."
        )
    _bullet_list(doc, recomendaciones_tecnicas)

    doc.add_heading("Propuesta de Conservación de Evidencias", level=1)
    doc.add_paragraph("Se recomienda conservar:")
    _bullet_list(doc, [
        "Copia verificada de la evidencia original, junto con su huella criptográfica (SHA-256).",
        "La base de datos de la cadena de custodia digital completa del caso.",
        "El fichero de modelo certificado (con su manifest de versión y hash) utilizado en el análisis.",
        "El presente informe pericial, en su versión final firmada.",
    ])
    doc.add_paragraph("Duración mínima de conservación recomendada: 5 años, conforme a estándares internacionales de referencia.")

    doc.add_page_break()

    # ============ DECLARACIONES FINALES ============
    doc.add_heading("Declaración de Objetividad", level=1)
    doc.add_paragraph(f"El abajo firmante, {perito_nombre}, declara:")
    _bullet_list(doc, [
        "Que ha actuado con objetividad, imparcialidad y rigor técnico en la elaboración del presente informe.",
        "Que no mantiene relación personal ni profesional con las partes implicadas.",
        "Que el informe se ha elaborado siguiendo criterios técnicos y científicos reconocidos.",
    ])

    doc.add_heading("Declaración de Independencia", level=1)
    doc.add_paragraph("El perito certifica:")
    _bullet_list(doc, [
        "No encontrarse en causa de recusación.",
        "No haber recibido instrucción alguna que comprometa la neutralidad del análisis.",
        "Mantener independencia absoluta en las conclusiones alcanzadas.",
    ])

    doc.add_heading("Posibilidad de Ratificación", level=1)
    doc.add_paragraph("El perito queda a disposición de la autoridad judicial para:")
    _bullet_list(doc, [
        "Ratificación presencial o telemática del presente informe.",
        "Aclaración técnica de cualquier extremo del análisis.",
        "Exhibición del material complementario (cadena de custodia completa, modelo certificado, código fuente del sistema).",
    ])

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Firma del perito:").bold = True
    doc.add_paragraph()
    doc.add_paragraph(perito_nombre)
    doc.add_paragraph(perito_titulacion)
    doc.add_paragraph(f"Colegiado n.º {perito_colegiado}")

    buffer = io.BytesIO()
    doc.save(buffer)
    doc_bytes = buffer.getvalue()

    if output_path:
        Path(output_path).write_bytes(doc_bytes)

    return doc_bytes
